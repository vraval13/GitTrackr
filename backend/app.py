from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from typing import Tuple, Dict, Optional, List
import re
import os
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import requests
import logging
import io
from flask import send_file
from datetime import datetime
import json

# ReportLab imports for enhanced PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, 
    Spacer, Image as RLImage, PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

app = Flask(__name__)
CORS(app, resources={r"/api/*": {"origins": "*"}})

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

GITHUB_API_TOKEN = os.getenv('GITHUB_API_TOKEN')
GITHUB_API_BASE = 'https://api.github.com'

# ==================== Utility Functions ====================

def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        raise ValueError("Failed to extract text from PDF")
    return text

def extract_text_and_links_from_pdf(file_path: str) -> Tuple[str, List[str]]:
    """Extract text and link annotations (URIs) from a PDF"""
    text = ""
    links: List[str] = []
    try:
        reader = PyPDF2.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
            try:
                annots = page.get("/Annots")
            except Exception:
                annots = None
            if annots:
                for annot in annots:
                    try:
                        obj = annot.get_object()
                        a = obj.get("/A") if obj else None
                        if a and a.get("/URI"):
                            uri = a.get("/URI")
                            if isinstance(uri, str):
                                links.append(uri)
                    except Exception:
                        continue
    except Exception as e:
        logger.error(f"Error extracting PDF (text+links): {e}")
        raise ValueError("Failed to extract text/links from PDF")
    return text, links

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        raise ValueError("Failed to extract text from DOCX")
    return text

def extract_text_and_links_from_docx(file_path: str) -> Tuple[str, List[str]]:
    """Extract text and external hyperlinks from DOCX"""
    text = ""
    links: List[str] = []
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        try:
            for rel in doc.part.rels.values():
                if rel.reltype and 'hyperlink' in rel.reltype:
                    target = getattr(rel, 'target_ref', None)
                    if target:
                        links.append(target)
        except Exception:
            pass
    except Exception as e:
        logger.error(f"Error extracting DOCX (text+links): {e}")
        raise ValueError("Failed to extract text/links from DOCX")
    return text, links

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting TXT: {e}")
        raise ValueError("Failed to extract text from TXT")

def extract_github_username(text: str, urls: Optional[List[str]] = None) -> Optional[str]:
    """Extract GitHub username/profile from text or provided URLs"""
    if not text and not urls:
        return None

    url_pattern = re.compile(
        r'(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9-]{1,39})(?=[/\s\)\]\.\,;:]|$)',
        re.IGNORECASE
    )
    if urls:
        for u in urls:
            if not u:
                continue
            m = url_pattern.search(u)
            if m:
                return m.group(1).strip().strip('.,;)')

    m = url_pattern.search(text or "")
    if m:
        return m.group(1).strip().strip('.,;)')

    label_pattern = re.compile(r'github[:\s\-]+([A-Za-z0-9-]{1,39})', re.IGNORECASE)
    m = label_pattern.search(text or "")
    if m:
        return m.group(1).strip().strip('.,;)')

    at_pattern = re.compile(r'(?<![\w\.-])@([A-Za-z0-9-]{1,39})(?!\.[A-Za-z0-9])')
    m = at_pattern.search(text or "")
    if m:
        return m.group(1).strip().strip('.,;)')

    return None

def validate_github_username(username: str) -> bool:
    """Validate if GitHub username exists"""
    try:
        headers = {'Authorization': f'token {GITHUB_API_TOKEN}'} if GITHUB_API_TOKEN else {}
        response = requests.get(f'{GITHUB_API_BASE}/users/{username}', headers=headers, timeout=5)
        return response.status_code == 200
    except Exception as e:
        logger.error(f"Error validating GitHub username: {e}")
        return False

def fetch_github_user_data(username: str) -> Dict:
    """Fetch user data from GitHub API"""
    try:
        headers = {'Authorization': f'token {GITHUB_API_TOKEN}'} if GITHUB_API_TOKEN else {}
        response = requests.get(f'{GITHUB_API_BASE}/users/{username}', headers=headers, timeout=5)
        
        if response.status_code != 200:
            raise ValueError(f"GitHub user not found: {username}")
        
        user_data = response.json()
        return {
            'login': user_data.get('login'),
            'name': user_data.get('name'),
            'bio': user_data.get('bio'),
            'followers': user_data.get('followers'),
            'following': user_data.get('following'),
            'public_repos': user_data.get('public_repos'),
            'avatar_url': user_data.get('avatar_url'),
        }
    except Exception as e:
        logger.error(f"Error fetching GitHub user data: {e}")
        raise

def fetch_github_repositories(username: str) -> list:
    """Fetch repositories from GitHub API"""
    try:
        headers = {'Authorization': f'token {GITHUB_API_TOKEN}'} if GITHUB_API_TOKEN else {}
        
        repositories = []
        page = 1
        per_page = 100
        
        while len(repositories) < 100:
            response = requests.get(
                f'{GITHUB_API_BASE}/users/{username}/repos',
                params={'page': page, 'per_page': per_page, 'sort': 'stars'},
                headers=headers,
                timeout=5
            )
            
            if response.status_code != 200:
                break
            
            repos = response.json()
            if not repos:
                break
            
            for repo in repos:
                if not repo.get('fork'):
                    repositories.append({
                        'name': repo.get('name'),
                        'description': repo.get('description'),
                        'url': repo.get('html_url'),
                        'stars': repo.get('stargazers_count', 0),
                        'forks': repo.get('forks_count', 0),
                        'language': repo.get('language'),
                    })
            
            page += 1
        
        return sorted(repositories, key=lambda x: x['stars'], reverse=True)
    except Exception as e:
        logger.error(f"Error fetching repositories: {e}")
        return []

def generate_pdf_summary(user: Dict, repositories: List[Dict], contribution_activity: Dict, language_distribution: List[Dict]) -> bytes:
    """
    Generate a comprehensive PDF summary matching the dashboard layout.
    Includes profile, stats, contributions, language distribution, and detailed repository information.
    """
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles matching dashboard colors
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=26,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=10,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'SubTitle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#64748b'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#334155'),
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold'
    )
    
    repo_name_style = ParagraphStyle(
        'RepoName',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=6,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#475569'),
        spaceAfter=6,
        leading=14
    )
    
    meta_style = ParagraphStyle(
        'MetaText',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#64748b'),
        spaceAfter=4
    )
    
    url_style = ParagraphStyle(
        'URLText',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=8
    )
    
    stat_value_style = ParagraphStyle(
        'StatValue',
        parent=styles['Normal'],
        fontSize=24,
        textColor=colors.HexColor('#4f46e5'),
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    stat_label_style = ParagraphStyle(
        'StatLabel',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#64748b'),
        alignment=TA_CENTER,
        spaceAfter=8
    )
    
    # ==================== HEADER ====================
    title = Paragraph("GitHub Profile Summary", title_style)
    elements.append(title)
    
    date_text = Paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        subtitle_style
    )
    elements.append(date_text)
    elements.append(Spacer(1, 0.2 * inch))
    
    # ==================== USER PROFILE SECTION ====================
    
    # Try to fetch avatar
    avatar_img = None
    if user.get('avatar_url'):
        try:
            response = requests.get(user['avatar_url'], timeout=5)
            if response.status_code == 200:
                img_buffer = io.BytesIO(response.content)
                avatar_img = RLImage(img_buffer, width=1*inch, height=1*inch)
        except Exception as e:
            logger.debug(f"Could not fetch avatar: {e}")
    
    # Profile section with border
    profile_data = []
    
    # Name and username
    name_text = f"<b><font size=14>{user.get('name', 'N/A')}</font></b>"
    profile_data.append([Paragraph(name_text, body_style)])
    
    username_text = f"<font color='#4f46e5'>@{user.get('login', 'N/A')}</font>"
    profile_data.append([Paragraph(username_text, body_style)])
    
    # Bio
    if user.get('bio'):
        profile_data.append([Spacer(1, 0.1*inch)])
        profile_data.append([Paragraph(user['bio'], body_style)])
    
    # Create profile layout table
    if avatar_img:
        profile_table_data = [[avatar_img, profile_data]]
        profile_table = Table(profile_table_data, colWidths=[1.3*inch, 5.2*inch])
        profile_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('ROUNDEDCORNERS', [10, 10, 10, 10]),
        ]))
        elements.append(profile_table)
    else:
        profile_content_table = Table(profile_data, colWidths=[6.5*inch])
        profile_content_table.setStyle(TableStyle([
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('RIGHTPADDING', (0, 0), (-1, -1), 15),
            ('TOPPADDING', (0, 0), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
            ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
            ('ROUNDEDCORNERS', [10, 10, 10, 10]),
        ]))
        elements.append(profile_content_table)
    
    elements.append(Spacer(1, 0.3 * inch))
    
    # ==================== PROFILE STATISTICS ====================
    stats_heading = Paragraph("Profile Statistics", heading_style)
    elements.append(stats_heading)
    
    stats_data = [
        [
            Paragraph("<b>📦</b><br/><b>Public Repos</b>", meta_style),
            Paragraph("<b>👥</b><br/><b>Followers</b>", meta_style),
            Paragraph("<b>🔗</b><br/><b>Following</b>", meta_style)
        ],
        [
            Paragraph(f"<b><font size=16>{user.get('public_repos', 0)}</font></b>", body_style),
            Paragraph(f"<b><font size=16>{user.get('followers', 0)}</font></b>", body_style),
            Paragraph(f"<b><font size=16>{user.get('following', 0)}</font></b>", body_style)
        ]
    ]
    
    stats_table = Table(stats_data, colWidths=[2.17*inch, 2.17*inch, 2.17*inch])
    stats_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e0e7ff')),
        ('BACKGROUND', (0, 1), (-1, 1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('ROUNDEDCORNERS', [10, 10, 10, 10]),
    ]))
    
    elements.append(stats_table)
    elements.append(Spacer(1, 0.3 * inch))
    
    # ==================== CONTRIBUTION ACTIVITY ====================
    contrib_heading = Paragraph("Contribution Activity", heading_style)
    elements.append(contrib_heading)
    
    total_contrib = contribution_activity.get('total', 0)
    current_streak = contribution_activity.get('current_streak', 0)
    longest_streak = contribution_activity.get('longest_streak', 0)
    
    contrib_data = [
        [
            Paragraph("📊<br/><b>Total Contributions</b>", meta_style),
            Paragraph("🔥<br/><b>Current Streak</b>", meta_style),
            Paragraph("🏆<br/><b>Longest Streak</b>", meta_style)
        ],
        [
            Paragraph(f"<font size=18 color='#4f46e5'><b>{total_contrib}</b></font>", body_style),
            Paragraph(f"<font size=18 color='#10b981'><b>{current_streak}</b></font><br/><font size=8>days</font>", body_style),
            Paragraph(f"<font size=18 color='#8b5cf6'><b>{longest_streak}</b></font><br/><font size=8>days</font>", body_style)
        ]
    ]
    
    contrib_table = Table(contrib_data, colWidths=[2.17*inch, 2.17*inch, 2.17*inch])
    contrib_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eff6ff')),
        ('BACKGROUND', (0, 1), (-1, 1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('ROUNDEDCORNERS', [10, 10, 10, 10]),
    ]))
    
    elements.append(contrib_table)
    elements.append(Spacer(1, 0.3 * inch))
    
    # ==================== LANGUAGE DISTRIBUTION ====================
    if language_distribution and len(language_distribution) > 0:
        lang_heading = Paragraph("Language Distribution", heading_style)
        elements.append(lang_heading)
        
        # Top 10 languages
        top_languages = language_distribution[:10]
        
        lang_rows = []
        for lang_stat in top_languages:
            lang_name = lang_stat.get('language', 'Unknown')
            percentage = lang_stat.get('percentage', 0)
            
            # Create a visual bar representation
            bar_width = percentage / 100.0 * 4.5  # Scale to fit in inches
            
            lang_row = [
                Paragraph(f"<b>{lang_name}</b>", body_style),
                Paragraph(f"{percentage}%", meta_style)
            ]
            lang_rows.append(lang_row)
        
        if lang_rows:
            lang_table = Table(lang_rows, colWidths=[4.5*inch, 2*inch])
            lang_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('LEFTPADDING', (0, 0), (-1, -1), 12),
                ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                ('LINEBELOW', (0, 0), (-1, -2), 0.5, colors.HexColor('#e2e8f0')),
                ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
                ('ROUNDEDCORNERS', [10, 10, 10, 10]),
            ]))
            elements.append(lang_table)
        
        elements.append(Spacer(1, 0.4 * inch))
    
    # ==================== REPOSITORIES ====================
    repos_heading = Paragraph("Repositories", heading_style)
    elements.append(repos_heading)
    
    if not repositories:
        no_repos = Paragraph("No repositories found.", body_style)
        elements.append(no_repos)
    else:
        # Repository summary
        total_stars = sum(repo.get('stars', 0) for repo in repositories)
        total_forks = sum(repo.get('forks', 0) for repo in repositories)
        languages = set(repo.get('language') for repo in repositories if repo.get('language'))
        
        summary_text = f"<b>Total:</b> {len(repositories)} repositories | " \
                      f"<b>⭐ Stars:</b> {total_stars} | " \
                      f"<b>🔱 Forks:</b> {total_forks} | " \
                      f"<b>💻 Languages:</b> {len(languages)}"
        
        summary = Paragraph(summary_text, meta_style)
        elements.append(summary)
        elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#e2e8f0'), spaceBefore=10, spaceAfter=15))
        
        # Individual repositories
        for idx, repo in enumerate(repositories, 1):
            # Repository name
            repo_name = Paragraph(
                f"<b>{idx}. {repo.get('name', 'Unnamed')}</b>",
                repo_name_style
            )
            elements.append(repo_name)
            
            # Metadata
            language = repo.get('language') or 'N/A'
            stars = repo.get('stars', 0)
            forks = repo.get('forks', 0)
            
            metadata_text = f"⭐ {stars} stars  •  🔱 {forks} forks  •  💻 {language}"
            metadata = Paragraph(metadata_text, meta_style)
            elements.append(metadata)
            
            # Description
            if repo.get('description'):
                desc = Paragraph(repo['description'], body_style)
                elements.append(desc)
            else:
                elements.append(Paragraph("<i>No description available</i>", meta_style))
            
            # URL
            if repo.get('url'):
                url_para = Paragraph(
                    f"<link href='{repo['url']}'><u>{repo['url']}</u></link>",
                    url_style
                )
                elements.append(url_para)
            
            # Separator
            if idx < len(repositories):
                elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#e2e8f0'), spaceBefore=8, spaceAfter=8))
            
            # Page break every 6 repos
            if idx % 6 == 0 and idx < len(repositories):
                elements.append(PageBreak())
                elements.append(Paragraph("Repositories (continued)", heading_style))
                elements.append(Spacer(1, 0.1 * inch))
    
    # ==================== FOOTER ====================
    elements.append(Spacer(1, 0.3 * inch))
    footer_line = HRFlowable(width="100%", thickness=1, color=colors.HexColor('#cbd5e1'))
    elements.append(footer_line)
    elements.append(Spacer(1, 0.1 * inch))
    
    footer_text = Paragraph(
        "Built with Next.js • React • GitHub API • Generated by GitHub Resume Analyzer",
        ParagraphStyle('Footer', fontSize=8, textColor=colors.HexColor('#94a3b8'), alignment=TA_CENTER)
    )
    elements.append(footer_text)
    
    # Build PDF
    doc.build(elements)
    
    buffer.seek(0)
    return buffer.read()

def fetch_github_contributions(username: str) -> Dict:
    """Fetch contribution calendar via GitHub GraphQL and compute total, current streak and longest streak."""
    if not GITHUB_API_TOKEN:
        logger.warning("GITHUB_API_TOKEN not set: contribution stats may be unavailable")
        return {"total": 0, "current_streak": 0, "longest_streak": 0, "days": []}

    graphql_url = "https://api.github.com/graphql"
    query = """
    query($login:String!) {
      user(login: $login) {
        contributionsCollection {
          contributionCalendar {
            totalContributions
            weeks {
              contributionDays {
                date
                contributionCount
              }
            }
          }
        }
      }
    }
    """
    try:
        headers = {'Authorization': f'bearer {GITHUB_API_TOKEN}', 'Content-Type': 'application/json'}
        resp = requests.post(graphql_url, json={"query": query, "variables": {"login": username}}, headers=headers, timeout=10)
        if resp.status_code != 200:
            logger.warning(f"GraphQL contributions fetch failed {resp.status_code}: {resp.text}")
            return {"total": 0, "current_streak": 0, "longest_streak": 0, "days": []}

        data = resp.json()
        weeks = data.get("data", {}).get("user", {}).get("contributionsCollection", {}).get("contributionCalendar", {}).get("weeks", [])
        total = data.get("data", {}).get("user", {}).get("contributionsCollection", {}).get("contributionCalendar", {}).get("totalContributions", 0)

        # Flatten days in chronological order
        days = []
        for w in weeks:
            for d in w.get("contributionDays", []):
                days.append({"date": d.get("date"), "count": d.get("contributionCount", 0)})

        # Compute longest streak and current streak
        longest = 0
        current = 0
        temp = 0
        # ensure days sorted by date
        days_sorted = sorted(days, key=lambda x: x["date"])
        for idx, day in enumerate(days_sorted):
            if day["count"] > 0:
                temp += 1
            else:
                if temp > longest:
                    longest = temp
                temp = 0
        if temp > longest:
            longest = temp

        # current streak: count from last day backwards while count>0
        current = 0
        for day in reversed(days_sorted):
            if day["count"] > 0:
                current += 1
            else:
                break

        return {"total": int(total), "current_streak": int(current), "longest_streak": int(longest), "days": days_sorted}
    except Exception as e:
        logger.error(f"Error fetching contributions for {username}: {e}")
        return {"total": 0, "current_streak": 0, "longest_streak": 0, "days": []}


def aggregate_language_distribution(username: str, repositories: List[Dict]) -> List[Dict]:
    """Aggregate language bytes across all repositories and return percentage distribution."""
    try:
        headers = {'Authorization': f'token {GITHUB_API_TOKEN}'} if GITHUB_API_TOKEN else {}
        lang_totals: Dict[str, int] = {}

        # For each repo, call languages endpoint to get bytes per language
        for repo in repositories:
            repo_name = repo.get("name")
            if not repo_name:
                continue
            try:
                resp = requests.get(f"{GITHUB_API_BASE}/repos/{username}/{repo_name}/languages", headers=headers, timeout=8)
                if resp.status_code != 200:
                    # fallback: use repository.language as a single-language count
                    primary = repo.get("language")
                    if primary:
                        lang_totals[primary] = lang_totals.get(primary, 0) + 1
                    continue
                langs = resp.json()
                for lang, bytes_count in langs.items():
                    lang_totals[lang] = lang_totals.get(lang, 0) + int(bytes_count)
            except Exception:
                primary = repo.get("language")
                if primary:
                    lang_totals[primary] = lang_totals.get(primary, 0) + 1
                continue

        total_bytes = sum(lang_totals.values()) or 1
        distribution = [{"language": k, "percentage": round(v * 100.0 / total_bytes, 1)} for k, v in lang_totals.items()]
        distribution.sort(key=lambda x: x["percentage"], reverse=True)
        return distribution
    except Exception as e:
        logger.error(f"Error aggregating languages for {username}: {e}")
        return []

# ==================== API Endpoints ====================

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'ok'})

@app.route('/api/upload', methods=['POST'])
def upload_resume():
    """Handle resume file upload and GitHub extraction"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type. Use PDF, DOCX or TXT'}), 400

        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        file_ext = filename.rsplit('.', 1)[1].lower()
        text = ""
        links: List[str] = []

        if file_ext == 'pdf':
            text, links = extract_text_and_links_from_pdf(file_path)
        elif file_ext == 'docx':
            text, links = extract_text_and_links_from_docx(file_path)
        elif file_ext == 'txt':
            text = extract_text_from_txt(file_path)
            links = []
        else:
            os.remove(file_path)
            return jsonify({'error': 'Unsupported file type'}), 400

        try:
            os.remove(file_path)
        except Exception:
            logger.debug("Could not remove uploaded file")

        github_username = extract_github_username(text, urls=links)

        if not github_username:
            return jsonify({'error': 'No GitHub profile found in resume'}), 404

        if not validate_github_username(github_username):
            return jsonify({'error': f'Invalid GitHub username: {github_username}'}), 404

        return jsonify({'github_username': github_username}), 200

    except Exception as e:
        logger.error(f"Error in upload_resume: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze_resume():
    """Handle resume text analysis and GitHub extraction"""
    try:
        data = request.get_json()
        
        if not data or 'text' not in data:
            return jsonify({'error': 'No text provided'}), 400
        
        text = data['text']
        
        if not text or not text.strip():
            return jsonify({'error': 'Resume text cannot be empty'}), 400
        
        github_username = extract_github_username(text)
        
        if not github_username:
            return jsonify({'error': 'No GitHub profile found in resume text'}), 404
        
        if not validate_github_username(github_username):
            return jsonify({'error': f'Invalid GitHub username: {github_username}'}), 404
        
        return jsonify({'github_username': github_username}), 200
    
    except Exception as e:
        logger.error(f"Error in analyze_resume: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/github/<username>', methods=['GET'])
def get_github_profile(username: str):
    try:
        # basic validation
        if not re.match(r'^[A-Za-z0-9_-]+$', username):
            return jsonify({'error': 'Invalid GitHub username'}), 400

        user_data = fetch_github_user_data(username)
        repositories = fetch_github_repositories(username)

        contribution_activity = fetch_github_contributions(username)
        language_distribution = aggregate_language_distribution(username, repositories)

        return jsonify({
            'user': user_data,
            'repositories': repositories,
            'contribution_activity': contribution_activity,
            'language_distribution': language_distribution
        }), 200
    except Exception as e:
        logger.error(f"Error in get_github_profile: {e}")
        return jsonify({'error': 'Failed to fetch GitHub profile'}), 500

@app.route('/api/github/<username>/export', methods=['GET'])
def export_github_profile(username: str):
    """Export comprehensive PDF summary of GitHub profile matching dashboard layout"""
    try:
        # Validate username format
        if not re.match(r'^[A-Za-z0-9_-]+$', username):
            return jsonify({'error': 'Invalid GitHub username'}), 400

        # Fetch user profile data
        user_data = fetch_github_user_data(username)
        
        # Fetch repositories
        repositories = fetch_github_repositories(username)
        
        # Fetch contribution activity (total, current streak, longest streak)
        contribution_activity = fetch_github_contributions(username)
        
        # Fetch language distribution across all repositories
        language_distribution = aggregate_language_distribution(username, repositories)

        # Generate comprehensive PDF with all dashboard features
        pdf_bytes = generate_pdf_summary(
            user=user_data,
            repositories=repositories,
            contribution_activity=contribution_activity,
            language_distribution=language_distribution
        )
        
        # Send PDF file as downloadable attachment
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'{username}_github_profile.pdf'
        )
        
    except ValueError as e:
        logger.warning(f"Export validation error for {username}: {e}")
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        logger.error(f"Error generating PDF for {username}: {e}")
        return jsonify({'error': 'Failed to generate PDF summary'}), 500
        
@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large error"""
    return jsonify({'error': 'File too large. Maximum size is 10MB'}), 413

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors"""
    return jsonify({'error': 'Endpoint not found'}), 404

# ==================== Main ====================

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)